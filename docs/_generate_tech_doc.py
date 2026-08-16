"""Generate DevSentinel in-depth technical documentation as a Word file."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

NAVY = RGBColor(0x0F, 0x17, 0x2A)
SLATE = RGBColor(0x33, 0x41, 0x55)
ACCENT = RGBColor(0x1D, 0x4E, 0xD8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "0F172A"
TABLE_ALT_BG = "F1F5F9"
CODE_BG = "F8FAFC"


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CBD5E1")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, bold=False, color=None, size=10, font="Calibri"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, name=font, size=size, bold=bold, color=color or SLATE)


def add_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DevSentinel Technical Documentation  ·  ")
    set_run_font(run, size=9, color=RGBColor(0x64, 0x74, 0x8B))
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r2 = p.add_run()
    r2._r.append(fld1)
    r2._r.append(instr)
    r2._r.append(fld2)
    set_run_font(r2, size=9, color=RGBColor(0x64, 0x74, 0x8B))


class DocBuilder:
    def __init__(self) -> None:
        self.doc = Document()
        self._setup()

    def _setup(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.2)
        add_page_number(section)

        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = SLATE
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        for i, size in ((1, 18), (2, 14), (3, 12)):
            s = styles[f"Heading {i}"]
            s.font.name = "Calibri"
            s.font.bold = True
            s.font.size = Pt(size)
            s.font.color.rgb = NAVY
            s.paragraph_format.space_before = Pt(16 if i == 1 else 12)
            s.paragraph_format.space_after = Pt(8)

    def h(self, level: int, text: str) -> None:
        self.doc.add_heading(text, level=level)

    def p(self, text: str, *, italic=False, bold=False, size=11) -> None:
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=SLATE)

    def bullets(self, items: list[str]) -> None:
        for item in items:
            para = self.doc.add_paragraph(style="List Bullet")
            para.clear()
            run = para.add_run(item)
            set_run_font(run, size=11, color=SLATE)

    def numbered(self, items: list[str]) -> None:
        for item in items:
            para = self.doc.add_paragraph(style="List Number")
            para.clear()
            run = para.add_run(item)
            set_run_font(run, size=11, color=SLATE)

    def code(self, text: str) -> None:
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(10)
        para.paragraph_format.left_indent = Cm(0.3)
        run = para.add_run(text)
        set_run_font(run, name="Consolas", size=9, color=NAVY)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), CODE_BG)
        shd.set(qn("w:val"), "clear")
        pPr.append(shd)

    def table(self, headers: list[str], rows: list[list[str]], col_widths=None) -> None:
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            set_cell_text(cell, h, bold=True, color=WHITE, size=9)
            shade_cell(cell, TABLE_HEADER_BG)
            set_cell_borders(cell)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = tbl.rows[r_idx + 1].cells[c_idx]
                set_cell_text(cell, val, size=9)
                if r_idx % 2 == 1:
                    shade_cell(cell, TABLE_ALT_BG)
                set_cell_borders(cell)
        if col_widths:
            for row in tbl.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Inches(w)
        self.doc.add_paragraph()

    def page_break(self) -> None:
        self.doc.add_page_break()

    def save(self, path: Path) -> None:
        self.doc.save(str(path))


def build() -> None:
    d = DocBuilder()

    # ── Cover ──────────────────────────────────────────────────────────────
    for _ in range(4):
        d.doc.add_paragraph()
    t = d.doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("DEVSENTINEL")
    set_run_font(r, size=14, bold=True, color=ACCENT)

    t = d.doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Technical Documentation")
    set_run_font(r, size=32, bold=True, color=NAVY)

    t = d.doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("In-depth architecture, data model, API, security, and operations reference")
    set_run_font(r, size=13, italic=True, color=SLATE)

    d.doc.add_paragraph()
    meta = [
        ("Document type", "Internal / engineering technical specification"),
        ("Product", "DevSentinel — AI-powered developer reliability platform"),
        ("Scope", "Full stack: Next.js frontend, FastAPI backend, PostgreSQL, Redis, integrations"),
        ("Version", "1.0  (codebase snapshot, August 2026)"),
        ("Classification", "Technical — not a customer marketing document"),
    ]
    d.table(["Field", "Value"], [[a, b] for a, b in meta])

    d.p(
        "This document describes every major technical aspect of the DevSentinel "
        "codebase as implemented: product behaviour, system architecture, database "
        "schema, authentication, APIs, AI pipelines, GitHub and Sentry integrations, "
        "real-time WebSockets, frontend structure, security controls, tests, and "
        "deployment. It is derived from the source in apps/web and apps/api, plus "
        "deployment manifests at the repository root."
    )

    d.page_break()

    # ── 1. Executive summary ───────────────────────────────────────────────
    d.h(1, "1. Executive summary")
    d.p(
        "DevSentinel is a multi-tenant SaaS that sits between pre-merge code review "
        "and post-deploy incident response. A GitHub App webhook delivers pull-request "
        "diffs to a FastAPI service, which asks Anthropic Claude to produce a structured "
        "review (0–100 score, summary, line-level findings). Low scores or critical "
        "findings automatically open an incident. Independently, a Sentry issue-alert "
        "webhook creates an AI-triaged Incident Room with live chat over WebSockets "
        "(Redis pub/sub, with an in-memory fallback). A Next.js 16 dashboard shows PR "
        "quality, team stats, branch assignments, weekly reports, and billing."
    )
    d.p("The product is organised around three loops:")
    d.bullets([
        "Pre-ship: GitHub PR opened/synchronized → installation token → unified diff → Claude review → PostgreSQL + GitHub review comment → optional auto-incident and email.",
        "Post-ship: Sentry alert → HMAC verify → Claude triage (root cause, suggested fix, P1–P4) → incident row → Redis/WebSocket broadcast → Incident Room.",
        "Team insights: per-member and per-repo aggregates, Claude Haiku quality grade, APScheduler weekly reports every Sunday 23:55 America/New_York.",
    ])
    d.p(
        "Auth is delegated to Supabase (email/password and Google OAuth). The API "
        "verifies JWTs (HS256 with SUPABASE_JWT_SECRET, or ES256 via JWKS). Org "
        "context is resolved from JWT app_metadata.org_id or the X-Org-Id header. "
        "Email verification is enforced in API middleware independently of the "
        "Supabase dashboard setting."
    )
    d.p(
        "Deployment target is free-tier: Vercel (frontend), Render (API), Supabase "
        "(Postgres + Auth), Upstash Redis, Resend for email, cron-job.org to keep "
        "the Render instance and in-process scheduler awake."
    )

    # ── 2. Product capabilities ─────────────────────────────────────────────
    d.h(1, "2. Product capabilities")

    d.h(2, "2.1 AI pull-request review")
    d.p(
        "When a GitHub App installation delivers pull_request events with action "
        "opened or synchronize, the API fetches the unified diff (Accept: "
        "application/vnd.github.v3.diff), truncates it to 500 lines, and sends it "
        "to Claude Sonnet with a JSON-only system prompt. The response is stored as "
        "a pull_requests row plus review_comments. Findings are posted back to GitHub "
        "as a single COMMENT review whose body includes a formatted findings block "
        "(GitHub inline positions are diff-relative, so file:line numbers are not "
        "sent as GitHub review comments to avoid HTTP 422)."
    )
    d.p("Scoring contract (0–100, not 0–10):")
    d.table(
        ["Score", "Meaning", "Downstream effect"],
        [
            ["100", "Flawless", "Green / info in dashboard"],
            ["80–99", "Good", "Info band"],
            ["60–79", "Needs work", "Warning band"],
            ["< 60", "Serious problems", "Critical band; auto-incident opened"],
            ["< 40", "Severe", "Auto-incident severity P1"],
        ],
    )
    d.p(
        "Auto-incident rule: score < 60 OR any comment with severity critical. "
        "P1 if score < 40, otherwise P2. The first critical comment body (or the "
        "summary) becomes root_cause. Affected files are the unique file paths of "
        "critical comments."
    )

    d.h(2, "2.2 Incident Room")
    d.p(
        "Incidents are created from Sentry webhooks, from critical PR reviews, or "
        "manually via POST /incidents. Status values: active, investigating, resolved. "
        "Severity on create is P0–P3 (API validation); Sentry triage returns P1–P4. "
        "Resolving sets resolved_at and mttr (minutes since created_at)."
    )
    d.p(
        "The Incident Room UI is a two-pane layout: live chat (ChatFeed) and a triage "
        "panel (root cause, suggested fix, affected files). Clients obtain a 5-minute "
        "HS256 JWT from GET /orgs/ws-token, then connect to "
        "ws://{API}/ws/incidents/{incident_id}?token=.... Dashboard live updates use "
        "the same endpoint with incident_id=dashboard."
    )

    d.h(2, "2.3 Organisations, members, invitations")
    d.p(
        "A verified user creates an org (unique slug) and becomes admin. Admins invite "
        "by email (role admin|member). The API records invitations and, if "
        "SUPABASE_URL + SUPABASE_SERVICE_KEY are set, calls Supabase Auth invite with "
        "redirect_to = {FRONTEND_URL}/join?org_id=.... Existing Supabase users (HTTP "
        "422 email_exists) still get a pending invitation visible in the dashboard "
        "banner. Join matches invitation email to JWT email. Last-admin and "
        "self-removal are blocked."
    )

    d.h(2, "2.4 GitHub App per organisation")
    d.p(
        "Each org can store its own GitHub App name, App ID, webhook secret, and "
        "private key (PEM normalised on save). Installation is linked via "
        "POST /orgs/github/link (installation_id). Repos are registered from "
        "installation webhooks, PR webhooks, and an admin sync endpoint that "
        "re-pulls GET /installation/repositories, adopts orphans filed under the "
        "wrong org, and backfills github_installation_id."
    )

    d.h(2, "2.5 Branch assignments and personal GitHub activity")
    d.p(
        "Admins assign any member to a branch; members may only self-assign. "
        "Uniqueness is (org, repo, branch, user). My GitHub activity matches "
        "user_profiles.github_login to pull_requests.author_github_login (case-"
        "insensitive for team stats). Branch activity lists PRs on head_branch, "
        "assigned engineers, and the caller’s recent commits from GitHub."
    )

    d.h(2, "2.6 Team quality and weekly reports")
    d.p(
        "GET /orgs/team-stats aggregates per-member PR counts, averages, critical/"
        "warning counts, riskiest file, per-repo averages, live branches (capped at "
        "20), and an optional Claude Haiku analysis (overallScore, grade A+–F, "
        "summary, strengths, risks, recommendation). Weekly reports persist the same "
        "JSON blob. Cron: Sunday 23:55 EST via APScheduler; admins can also POST "
        "/orgs/weekly-report/generate (10/hour)."
    )

    d.h(2, "2.7 Notifications")
    d.p(
        "Email channels (Resend) subscribe to incident_created and/or "
        "pr_review_completed. Config JSON is {emails: [...]}. Admins can test a "
        "channel. If RESEND_API_KEY is unset, sends are skipped with a warning."
    )

    d.h(2, "2.8 Billing (UI only)")
    d.p(
        "Plans in types: free | pro | team. Stripe keys exist in Settings and "
        "render.yaml, and @stripe/stripe-js is a frontend dependency, but the "
        "billing settings page is currently a static mock (Free $0 / Pro $29 / "
        "Team $79, sample invoices). There is no Stripe webhook router in the API. "
        "organizations.plan defaults to free; stripe_customer_id is unused in "
        "application logic."
    )

    d.h(2, "2.9 Marketing surface")
    d.p(
        "Public App Router group (marketing) includes landing, features, pricing, "
        "docs, API reference (static endpoint catalogue), about, blog, changelog, "
        "roadmap, careers, press, support, status, and coming-soon. Authenticated "
        "users hitting /, /login, or /sign-up are redirected to /dashboard."
    )

    # ── 3. Architecture ────────────────────────────────────────────────────
    d.h(1, "3. System architecture")

    d.h(2, "3.1 Logical topology")
    d.code(
        "Browser (Next.js 16 / React 19)\n"
        "  |  HTTPS REST  +  WSS\n"
        "  v\n"
        "Vercel  ──►  Render (FastAPI / uvicorn)\n"
        "               |         |         |          |\n"
        "               v         v         v          v\n"
        "         Supabase     Upstash    Anthropic   GitHub App API\n"
        "         Postgres     Redis      Claude      (install tokens,\n"
        "         + Auth       pub/sub    Sonnet/     diffs, reviews,\n"
        "                                  Haiku       branches, commits)\n"
        "               ^\n"
        "         Sentry issue alerts  (HMAC optional)\n"
        "         Resend transactional email\n"
        "         cron-job.org  GET /health  (~10 min keep-warm)"
    )

    d.h(2, "3.2 Request paths")
    d.table(
        ["Path", "Entry", "Auth", "Persistence / side effects"],
        [
            ["User REST", "Bearer JWT + X-Org-Id", "Supabase JWT + email verified", "SQLAlchemy async → Postgres"],
            ["GitHub webhook", "POST /webhooks/github", "HMAC-SHA256 (per-org or global secret)", "PR + comments; GitHub review; Redis; email"],
            ["Sentry webhook", "POST /webhooks/sentry?org_id=", "Optional HMAC; org_id is query param", "Incident + Redis + email"],
            ["WebSocket", "GET /ws/incidents/{id}?token=", "Short-lived HS256 JWT_SECRET", "Messages / resolve + pub/sub"],
            ["OAuth callback", "GET /auth/callback", "Supabase PKCE code", "Upsert /users/profile"],
            ["GitHub install callback", "GET /api/github/callback", "Supabase session", "POST /orgs/github/link"],
        ],
    )

    d.h(2, "3.3 Multi-tenancy model")
    d.p(
        "Tenancy is organisation-scoped, not row-level security in Postgres. Almost "
        "every authenticated query filters by org_id from get_verified_org_id. "
        "Members bind a Supabase user UUID (JWT sub) to an org with a role. Repos, "
        "PRs, incidents, invitations, branch assignments, weekly reports, and "
        "notification channels all hang off organizations.id."
    )
    d.p(
        "Webhook org resolution is a known multi-tenant edge: if installation_id is "
        "not stored on any org, the code falls back to an existing repo with that "
        "installation, then to the first organisation in the database (single-tenant "
        "dev fallback). POST /orgs/github/sync-repos exists specifically to heal "
        "mis-filed repos and persist installation_id."
    )

    d.h(2, "3.4 Process model")
    d.p(
        "The API is a single uvicorn process on Render free tier. SQLAlchemy uses "
        "NullPool (no persistent pool) with statement_cache_size=0 so PgBouncer "
        "transaction mode is less likely to break prepared statements — but the "
        "documented requirement is still the session pooler on port 5432, not 6543. "
        "SSL is enabled with hostname verification disabled (common for pooler "
        "hostnames). APScheduler runs in-process; if the instance sleeps, Sunday "
        "reports do not fire unless keep-warm pings succeed."
    )
    d.p(
        "WebSocket connections are held in a process-local dict org_id → set[WebSocket]. "
        "Redis pub/sub is the multi-process fan-out. If Redis is down, publish_to_org "
        "calls broadcast_to_org in the same process only."
    )

    # ── 4. Tech stack ──────────────────────────────────────────────────────
    d.h(1, "4. Technology stack")
    d.h(2, "4.1 Frontend (apps/web)")
    d.table(
        ["Component", "Choice", "Notes"],
        [
            ["Framework", "Next.js 16.2.6 App Router", "Turbopack root pinned to apps/web"],
            ["UI library", "React 19.2.4", "Client layouts for app shell"],
            ["Language", "TypeScript 5", "strict project tsconfig"],
            ["Styling", "Tailwind CSS v4", "PostCSS; CSS variables for brand tokens"],
            ["Motion", "Framer Motion 12", "Onboarding, marketing hero"],
            ["Components", "shadcn / Base UI / Radix", "Button, Avatar, Slot"],
            ["Data fetching", "SWR 2.4", "30s PR/incident poll; 60s team/weekly; paused when tab hidden"],
            ["Auth client", "@supabase/ssr + supabase-js", "Browser and server cookie clients"],
            ["Payments UI", "stripe + @stripe/stripe-js", "Present; checkout not wired"],
            ["Tests", "Vitest 2", "password.ts unit tests"],
            ["Lint", "ESLint 9 + eslint-config-next", "npm run lint"],
        ],
    )

    d.h(2, "4.2 Backend (apps/api)")
    d.table(
        ["Component", "Choice", "Notes"],
        [
            ["Framework", "FastAPI 0.115 + uvicorn 0.30", "ASGI; /health for probes"],
            ["Language", "Python 3.11+ (Render 3.12)", "async throughout"],
            ["ORM", "SQLAlchemy 2.x asyncio", "Mapped / mapped_column; create_all on startup"],
            ["Migrations", "Alembic 1.13", "Revisions 001–007; async env.py"],
            ["Driver", "asyncpg 0.29", "postgresql+asyncpg:// URLs"],
            ["Auth JWT", "python-jose[cryptography]", "HS256 Supabase + RS256 GitHub App JWT"],
            ["HTTP out", "httpx 0.27", "GitHub, Supabase invite, Resend, JWKS"],
            ["AI", "anthropic 0.34", "AsyncAnthropic singleton"],
            ["Cache / pubsub", "redis[asyncio] 5.0", "Optional; ping cached"],
            ["Rate limit", "slowapi 0.1.9", "Per-route; memory or Redis storage"],
            ["Scheduler", "APScheduler 3.10 + pytz", "ImportError disables cron"],
            ["Config", "pydantic-settings 2.4", "Settings in models/database.py"],
            ["Tests", "pytest 8 + pytest-asyncio", "asyncio_mode = auto"],
        ],
    )

    d.h(2, "4.3 External platforms")
    d.table(
        ["Platform", "Role"],
        [
            ["Supabase", "Postgres (session pooler :5432), Auth (email confirm, Google OAuth, invite API), JWT issuer"],
            ["Anthropic", "claude-sonnet-4-6 for PR review and incident triage; claude-haiku-4-5-20251001 for team grade"],
            ["GitHub Apps", "Webhooks, installation tokens, diffs, reviews, repo/branch/commit lists"],
            ["Sentry", "Issue-created webhooks into /webhooks/sentry?org_id="],
            ["Resend", "Transactional HTML email for channels"],
            ["Upstash Redis", "rediss:// pub/sub in production"],
            ["Vercel", "Next.js hosting; security headers from next.config.ts"],
            ["Render", "Blueprint render.yaml; free web service; env sync:false"],
            ["Stripe", "Keys reserved; no live billing pipeline yet"],
        ],
    )

    # ── 5. Repository ──────────────────────────────────────────────────────
    d.h(1, "5. Repository structure")
    d.code(
        "devsentinel/\n"
        "├── apps/\n"
        "│   ├── web/                      Next.js application\n"
        "│   │   ├── app/\n"
        "│   │   │   ├── (app)/            Authenticated product (dashboard, incidents, settings)\n"
        "│   │   │   ├── (marketing)/      Public pages\n"
        "│   │   │   ├── api/github/callback/   GitHub App install return\n"
        "│   │   │   ├── auth/callback/    Supabase OAuth / magic-link exchange\n"
        "│   │   │   ├── login, sign-up, pricing, page.tsx\n"
        "│   │   ├── components/           UI, auth, dashboard, incidents, settings, layout\n"
        "│   │   ├── contexts/org-context.tsx\n"
        "│   │   ├── hooks/                use-api, WebSocket, animation helpers\n"
        "│   │   ├── lib/                  api.ts, supabase clients, password, utils\n"
        "│   │   ├── types/index.ts        Shared domain types\n"
        "│   │   ├── middleware.ts         Auth gate + cookie refresh\n"
        "│   │   └── next.config.ts        CSP / HSTS / X-Frame-Options\n"
        "│   └── api/                      FastAPI application\n"
        "│       ├── main.py               App factory, CORS, routers, scheduler, create_all\n"
        "│       ├── models/               SQLAlchemy entities + Settings + engine\n"
        "│       ├── routers/              webhooks, prs, incidents, orgs, users, notifications, ws\n"
        "│       ├── services/             claude, github, redis, email, weekly report\n"
        "│       ├── middleware/           auth.py, security.py\n"
        "│       ├── alembic/versions/     001–007\n"
        "│       └── tests/\n"
        "├── render.yaml                   Render Blueprint\n"
        "├── vercel.json                   { \"framework\": \"nextjs\" }\n"
        "├── README.md\n"
        "└── docs/                         This document + design specs/plans"
    )
    d.p(
        "There is no npm/pnpm workspace root package; frontend and backend are "
        "independent apps. Vercel must be configured with Root Directory apps/web "
        "(or equivalent). Render uses rootDir: apps/api."
    )

    # ── 6. Data model ──────────────────────────────────────────────────────
    d.h(1, "6. Data model")
    d.p(
        "Primary keys are string UUIDs generated in Python (uuid.uuid4) except "
        "user_profiles.id, which is the Supabase user UUID. Timestamps use naive "
        "UTC datetime.utcnow. JSON payloads that are not first-class columns are "
        "stored as Text (affected_files, report_data, channel config/events)."
    )
    d.p(
        "Startup calls Base.metadata.create_all in addition to Alembic. New "
        "environments get tables from models; schema changes after that require "
        "alembic upgrade head. Operators should not rely on create_all for "
        "column additions."
    )

    d.h(2, "6.1 organizations")
    d.table(
        ["Column", "Type", "Notes"],
        [
            ["id", "String PK", "UUID"],
            ["name", "String(120)", "Display name"],
            ["slug", "String(80)", "Unique, indexed"],
            ["plan", "String(20)", "Default free"],
            ["stripe_customer_id", "String nullable", "Reserved"],
            ["created_at", "DateTime", "utcnow"],
            ["github_app_name", "String(120)", "Per-org GitHub App"],
            ["github_app_id", "String(40)", "Issuer for RS256 App JWT"],
            ["github_webhook_secret", "String/Text", "HMAC secret; never returned by GET /github/config"],
            ["github_private_key", "String/Text", "PEM; normalize_pem on save"],
            ["github_installation_id", "Integer", "Linked installation"],
        ],
    )

    d.h(2, "6.2 members")
    d.table(
        ["Column", "Type", "Notes"],
        [
            ["id", "String PK", "UUID"],
            ["org_id", "FK organizations CASCADE", "Tenant"],
            ["user_id", "String indexed", "Supabase sub"],
            ["name", "String(120)", "From metadata or email local-part"],
            ["email", "String(255)", "Invite matching key"],
            ["role", "String(20)", "admin | member"],
            ["joined_at", "DateTime", ""],
        ],
    )

    d.h(2, "6.3 repos")
    d.table(
        ["Column", "Type", "Notes"],
        [
            ["id", "String PK", "UUID"],
            ["org_id", "FK organizations CASCADE", ""],
            ["github_repo_id", "Integer", "GitHub numeric id"],
            ["name", "String(255)", "Short name"],
            ["full_name", "String(255)", "owner/repo"],
            ["installation_id", "Integer", "0 used as placeholder in webhook path"],
            ["is_active", "Boolean", "Monitoring toggle"],
        ],
    )

    d.h(2, "6.4 invitations")
    d.table(
        ["Column", "Type", "Notes"],
        [
            ["id", "String PK", "UUID"],
            ["org_id", "FK CASCADE", ""],
            ["email", "String(255) indexed", "Pending match"],
            ["role", "String(20)", "admin | member"],
            ["invited_by", "String", "Admin user_id"],
            ["status", "String(20)", "pending | accepted | declined"],
            ["created_at / accepted_at", "DateTime", "accepted_at also set on decline"],
        ],
    )

    d.h(2, "6.5 user_profiles")
    d.table(
        ["Column", "Type", "Notes"],
        [
            ["id", "String PK", "Supabase user UUID, not auto-generated"],
            ["email", "String(255) unique", ""],
            ["full_name", "String(255)", ""],
            ["github_login", "String(120)", "Stripped of leading @; links PR authorship"],
            ["created_at / updated_at", "DateTime", "onupdate utcnow"],
        ],
    )

    d.h(2, "6.6 pull_requests and review_comments")
    d.table(
        ["Column", "Type", "Notes"],
        [
            ["pull_requests.id", "String PK", ""],
            ["org_id", "String indexed", "Copied from repo.org_id"],
            ["repo_id", "FK repos CASCADE", ""],
            ["github_pr_number", "Integer", ""],
            ["title", "String(500)", ""],
            ["author_github_login", "String(120)", "GitHub login, not Supabase id"],
            ["head_branch", "String(255)", "PR source ref; indexed with repo_id"],
            ["status", "String(20)", "pending | reviewed | merged | closed (writers currently set reviewed)"],
            ["review_score", "Integer", "0–100"],
            ["summary", "String(2000)", "Claude summary"],
            ["created_at / updated_at", "DateTime", ""],
            ["review_comments.*", "", "file_path, line_number, severity critical|warning|info, body"],
        ],
    )

    d.h(2, "6.7 incidents and incident_messages")
    d.table(
        ["Column", "Type", "Notes"],
        [
            ["incidents.id", "String PK", ""],
            ["org_id", "String indexed", "No FK — filter-only tenancy"],
            ["repo_id", "String nullable", "Often unset for Sentry/manual"],
            ["sentry_issue_id", "String nullable", ""],
            ["title", "String(500)", ""],
            ["severity", "String(5)", "P1–P4 typical; create allows P0"],
            ["status", "String(20)", "active | investigating | resolved"],
            ["root_cause / suggested_fix", "Text", "AI or human"],
            ["affected_files", "Text", "JSON array string"],
            ["users_affected / error_rate", "Integer / Float", "Schema present; Sentry path does not populate"],
            ["mttr", "Integer", "Minutes; set on resolve"],
            ["resolved_at / created_at", "DateTime", ""],
            ["incident_messages", "", "user_id, author_name, body, is_ai, created_at"],
        ],
    )

    d.h(2, "6.8 branch_assignments")
    d.p("Unique index on (org_id, repo_id, branch_name, user_id). created_by stores the admin or self-assigning member.")

    d.h(2, "6.9 weekly_reports")
    d.p(
        "week_of is a Date (intended as the Sunday covered). report_data is a JSON "
        "serialisation of TeamStats (members, repos, orgStats, aiAnalysis). GET "
        "/orgs/weekly-report returns the newest by generated_at."
    )

    d.h(2, "6.10 notification_channels")
    d.p(
        "channel_type defaults to email. config and events are JSON text. Valid "
        "events: incident_created, pr_review_completed. is_enabled gates dispatch."
    )

    d.h(2, "6.11 Entity relationships")
    d.code(
        "Organization 1──* Member\n"
        "             1──* Repo 1──* PullRequest 1──* ReviewComment\n"
        "             1──* Invitation\n"
        "             1──* BranchAssignment  → Repo, user_id\n"
        "             1──* WeeklyReport\n"
        "             1──* NotificationChannel\n"
        "Incident (org_id) 1──* IncidentMessage\n"
        "UserProfile (id = Supabase sub)  ←  Member.user_id  (logical, no FK)"
    )

    d.h(2, "6.12 Alembic history")
    d.table(
        ["Rev", "Change"],
        [
            ["001", "invitations table + email index"],
            ["002", "user_profiles"],
            ["003", "GitHub App columns on organizations"],
            ["004", "branch_assignments + user_profiles.github_login"],
            ["005", "pull_requests.head_branch + (repo_id, head_branch) index"],
            ["006", "weekly_reports"],
            ["007", "notification_channels"],
        ],
    )
    d.p(
        "Core tables organizations, members, repos, pull_requests, review_comments, "
        "incidents, incident_messages are created by SQLAlchemy metadata / initial "
        "bootstrap rather than an Alembic 000 revision in this repo."
    )

    # ── 7. Auth ────────────────────────────────────────────────────────────
    d.h(1, "7. Authentication and authorisation")

    d.h(2, "7.1 Frontend session")
    d.p(
        "middleware.ts builds a Supabase server client from request cookies, calls "
        "getUser(), refreshes the session via setAll on cookies, and:"
    )
    d.bullets([
        "Redirects authenticated users from /, /login, /sign-up to /dashboard.",
        "Redirects unauthenticated users on non-public paths to /login.",
        "Treats as public: /, /pricing, /login, /sign-up, /api/webhooks, /auth/callback, and static assets excluded by the matcher.",
        "Marketing routes under (marketing) are not listed in PUBLIC_PATHS; they are still reachable because the matcher excludes many static files, but HTML marketing pages that are not in PUBLIC_PATHS will require a session. Operators should treat this as a middleware configuration detail when adding public pages.",
    ])
    d.p(
        "AuthProvider listens to onAuthStateChange. Login uses signInWithPassword then "
        "GET /orgs/mine (and /orgs/my-invites). Google OAuth redirectTo is "
        "{origin}/auth/callback. The callback exchanges the code, POSTs /users/profile "
        "(non-fatal on failure), and redirects to next (open-redirect guarded: must "
        "start with / but not //)."
    )
    d.p(
        "Password rule (lib/password.ts): length ≥ 8, lowercase, uppercase, digit. "
        "Score 0–4 also rewards symbols or length ≥ 12. Strength meter clamps score "
        "to 2 when invalid. Mirror this in the Supabase Auth password policy."
    )

    d.h(2, "7.2 API JWT verification")
    d.p("verify_supabase_token (HTTPBearer):")
    d.numbered([
        "Read unverified header alg (default HS256).",
        "HS256: decode with SUPABASE_JWT_SECRET, audience authenticated.",
        "Otherwise: read iss, fetch {iss}/.well-known/jwks.json (cached per issuer), match kid, retry once after cache clear on rotation, decode with that alg and audience authenticated.",
        "JWTError → 401 Invalid or expired token.",
    ])
    d.p(
        "require_verified_email: if ENFORCE_EMAIL_VERIFICATION is true (default), "
        "require user_metadata.email_verified == true OR payload.email_confirmed_at. "
        "OAuth users typically have the flag. Failure is 403 Email not verified. "
        "Safety valve: set ENFORCE_EMAIL_VERIFICATION=false."
    )
    d.p(
        "get_verified_org_id: org_id = app_metadata.org_id OR X-Org-Id header. "
        "Missing both → 401 No org context in token. The frontend stores org id in "
        "localStorage key devsentinel_org_id and sends X-Org-Id on every apiFetch. "
        "JWT app_metadata.org_id is not written by this codebase during org create, "
        "so production traffic depends on the header."
    )

    d.h(2, "7.3 Role checks")
    d.p(
        "Admin-only operations load Member for (org_id, sub) and require role == "
        "admin. Applied to: list members, update org, invite, remove member, cancel "
        "invite, save GitHub config, sync/link GitHub, toggle repo, weekly report "
        "generate, notification channel mutations. Branch assign: admins any user; "
        "members only themselves. GET /orgs/mine, /orgs/my-invites, POST /orgs/join "
        "and /orgs/decline do not require X-Org-Id."
    )

    d.h(2, "7.4 WebSocket tokens")
    d.p(
        "GET /orgs/ws-token issues HS256 JWT with JWT_SECRET, claims org_id, sub, "
        "name, iat, exp = now+300. Missing JWT_SECRET → 503. WS handler closes with "
        "code 4001 if decode fails. Name used as chat author_name."
    )

    # ── 8. API ─────────────────────────────────────────────────────────────
    d.h(1, "8. HTTP and WebSocket API")
    d.p(
        "Unless noted, JSON envelopes are {success: true, data: ...}. apiFetch "
        "unwraps that envelope. Notification channel routes return {data: ...} "
        "without success. Errors are FastAPI HTTPException detail strings."
    )

    d.h(2, "8.1 Health")
    d.table(
        ["Method", "Path", "Auth", "Behaviour"],
        [["GET", "/health", "None", '{"status": "ok"} — Render healthCheckPath and keep-warm']],
    )

    d.h(2, "8.2 Users")
    d.table(
        ["Method", "Path", "Auth / limit", "Behaviour"],
        [
            ["POST", "/users/profile", "Verified email; 30/hour", "Upsert profile by JWT sub/email"],
            ["GET", "/users/profile", "Bearer", "404 if missing"],
        ],
    )

    d.h(2, "8.3 Organisations")
    d.table(
        ["Method", "Path", "Auth / limit", "Behaviour"],
        [
            ["POST", "/orgs", "Verified email; 10/hour", "Create org + admin member; 409 on slug"],
            ["GET", "/orgs/me", "Org context", "Current org"],
            ["GET", "/orgs/mine", "Bearer only", "All memberships with role"],
            ["GET", "/orgs/ws-token", "Org context", "5-minute WS JWT"],
            ["GET", "/orgs/members", "Admin", "Members + invitations (all statuses)"],
            ["GET", "/orgs/my-invites", "Bearer", "Pending invites for JWT email"],
            ["POST", "/orgs/decline", "Bearer", "pending → declined"],
            ["PATCH", "/orgs", "Admin", "Name/slug; 409 on slug clash"],
            ["POST", "/orgs/invite", "Admin; 30/hour", "Invitation + optional Supabase invite"],
            ["DELETE", "/orgs/members/{id}", "Admin; 204", "No self; no last admin"],
            ["DELETE", "/orgs/invitations/{id}", "Admin; 204", "Pending only"],
            ["POST", "/orgs/join", "Bearer", "Accept invite; idempotent if already member"],
        ],
    )

    d.h(2, "8.4 GitHub integration and repos")
    d.table(
        ["Method", "Path", "Auth / limit", "Behaviour"],
        [
            ["GET", "/orgs/github/config", "Org", "isConfigured, appName, isConnected, installationId — no secrets"],
            ["POST", "/orgs/github/config", "Admin", "Save App credentials; PEM repaired"],
            ["GET", "/orgs/repos", "Org", "Monitored repos"],
            ["POST", "/orgs/github/sync-repos", "Admin; 20/hour", "Reconcile GitHub list; adopt orphans; backfill installation_id"],
            ["POST", "/orgs/github/link", "Admin", "Persist installation_id; reassign repos; backfill from GitHub API"],
            ["PATCH", "/orgs/repos/{id}", "Admin", "is_active toggle"],
            ["GET", "/orgs/repos/{id}/branches", "Org", "Live GitHub branches"],
            ["POST", "/orgs/repos/{id}/branches/assign", "Member/admin", "201; 409 duplicate"],
            ["DELETE", "/orgs/repos/{id}/branches/assign/{aid}", "Owner or admin; 204", "Remove assignment"],
            ["GET", "/orgs/branch-assignments", "Org", "Joined member + repo names"],
            ["GET", "/orgs/me/github-activity", "Org", "Caller PRs, assignments, stats"],
            ["PATCH", "/orgs/me/github-login", "Bearer", "Create profile if needed"],
            ["GET", "/orgs/branch-activity", "Org", "Query repo_id + branch"],
            ["GET", "/orgs/team-stats", "Org", "Aggregates + Haiku analysis"],
            ["GET", "/orgs/weekly-report", "Org", "Latest or null"],
            ["GET", "/orgs/weekly-reports", "Org", "All, newest first"],
            ["POST", "/orgs/weekly-report/generate", "Admin; 10/hour", "201 stored report"],
        ],
    )

    d.h(2, "8.5 Pull requests")
    d.table(
        ["Method", "Path", "Auth", "Behaviour"],
        [
            ["GET", "/prs", "Org", "Newest first; repo join for full_name"],
            ["GET", "/prs/{pr_id}", "Org", "Includes comments; 404 if wrong org"],
        ],
    )

    d.h(2, "8.6 Incidents")
    d.table(
        ["Method", "Path", "Auth", "Behaviour"],
        [
            ["GET", "/incidents", "Org", "Newest first, no messages"],
            ["GET", "/incidents/{id}", "Org", "With chronological messages"],
            ["POST", "/incidents", "Org; 201", "Manual create; severity P0–P3; emails queued"],
            ["PATCH", "/incidents/{id}", "Org", "status/severity/root_cause/suggested_fix; resolve computes MTTR"],
        ],
    )

    d.h(2, "8.7 Notifications")
    d.table(
        ["Method", "Path", "Auth", "Behaviour"],
        [
            ["GET", "/notifications/channels", "Org", "List"],
            ["POST", "/notifications/channels", "Admin; 201", "Validate events"],
            ["PATCH", "/notifications/channels/{id}", "Admin", "Partial update"],
            ["DELETE", "/notifications/channels/{id}", "Admin; 204", ""],
            ["POST", "/notifications/channels/{id}/test", "Admin", "Queue test email"],
        ],
    )

    d.h(2, "8.8 Webhooks")
    d.table(
        ["Method", "Path", "Auth", "Behaviour"],
        [
            ["POST", "/webhooks/github", "HMAC sha256=", "installation*, pull_request opened/synchronize"],
            ["POST", "/webhooks/sentry", "Optional HMAC; query org_id required", "action=created only"],
        ],
    )

    d.h(2, "8.9 WebSocket protocol")
    d.p("URL: /ws/incidents/{incident_id}?token={ws_jwt}")
    d.table(
        ["Direction", "type", "Payload / effect"],
        [
            ["Client → server", "message.send", "{body} persisted as IncidentMessage; published as message.new"],
            ["Client → server", "incident.resolve", "status resolved, mttr; published incident.resolved"],
            ["Server → client", "incident.new", "New incident (Sentry/PR); dashboard should listen"],
            ["Server → client", "message.new", "Chat line including authorInitials"],
            ["Server → client", "incident.resolved", "{id, resolvedAt, mttr}"],
        ],
    )
    d.p(
        "Frontend WSEvent also types incident.created, incident.updated, and "
        "pr.reviewed. The webhook publisher currently emits incident.new (not "
        "incident.created). Dashboard hook listens for incident.created — live "
        "dashboard incident toasts may miss webhook-created incidents unless the "
        "event names are aligned. SWR 30s polling still refreshes lists."
    )

    d.h(2, "8.10 CORS")
    d.p(
        "CORS_ORIGINS comma-separated; credentials allowed; methods and headers *. "
        "Default http://localhost:3000. Production must include the Vercel origin "
        "or browsers will block REST and the WS handshake from the app origin."
    )

    # ── 9. AI ──────────────────────────────────────────────────────────────
    d.h(1, "9. AI / LLM integration")
    d.p(
        "All calls live in apps/api/services/claude_service.py. One Anthropic "
        "async client is lazily constructed. There are no tools, agents, or multi-"
        "step graphs. Output is stripped of markdown fences then json.loads. A "
        "malformed model response raises and is handled per caller."
    )
    d.table(
        ["Function", "Model", "max_tokens", "Consumers"],
        [
            ["review_pull_request", "claude-sonnet-4-6", "2048", "GitHub webhook; score/comments/summary"],
            ["triage_incident", "claude-sonnet-4-6", "1024", "Sentry webhook; fallback dict on exception"],
            ["analyze_team_quality", "claude-haiku-4-5-20251001", "512", "team-stats + weekly reports; skipped if no PRs"],
        ],
    )

    d.h(2, "9.1 PR review prompt contract")
    d.p(
        "System: senior engineer; bugs, OWASP Top 10, performance; JSON only. "
        "User: repo, title, truncated diff. JSON: comments[{file, line, "
        "severity: critical|warning|info, body}], score integer 0–100, summary. "
        "Diff truncation: first 500 lines plus a truncated-N-lines trailer."
    )

    d.h(2, "9.2 Incident triage prompt contract")
    d.p(
        "System: senior reliability engineer; JSON only. User: title, stack trace "
        "(last 8 frames formatted), affected files, blame map (currently {} from "
        "Sentry handler). JSON: rootCause, suggestedFix, affectedFiles[], "
        "blastRadius, severity P1–P4. blastRadius is not persisted on Incident."
    )

    d.h(2, "9.3 Team quality prompt contract")
    d.p(
        "System: engineering manager; JSON only. User: repo count, total PRs, "
        "average score /100, critical/warning totals, per-engineer lines. JSON: "
        "overallScore 0–100, grade A+|A|A-|B+|B|B-|C+|C|C-|D|F, summary, "
        "strengths[≤3], risks[≤3], recommendation."
    )

    d.h(2, "9.4 Eval harness (designed, not fully productised)")
    d.p(
        "docs/superpowers/specs/2026-08-04-llm-eval-harness-design.md specifies a "
        "harness because unit tests mock get_client and never hit the model. "
        "Failure classes: broken JSON contract, quality regression, poor "
        "precision/recall, token cost. LangChain/LangGraph/promptfoo were rejected "
        "to avoid rewriting working single-shot calls and leaking diffs to extra "
        "vendors. Implementers should treat that spec as the source of truth for "
        "eval-driven changes to prompts."
    )

    # ── 10. GitHub ─────────────────────────────────────────────────────────
    d.h(1, "10. GitHub App integration")

    d.h(2, "10.1 Authentication to GitHub")
    d.p(
        "_get_app_jwt builds RS256 JWT: iat = now-60, exp = now+600, iss = App ID. "
        "Key resolution: explicit private_key argument → GITHUB_APP_PRIVATE_KEY → "
        "file at GITHUB_APP_PRIVATE_KEY_PATH (default ./github-app.pem). "
        "normalize_pem repairs collapsed newlines, missing armor, and 64-char wrap; "
        "it prefers a variant cryptography can parse (typically PKCS#1 RSA PRIVATE KEY)."
    )
    d.p(
        "Installation token: POST /app/installations/{id}/access_tokens with the App "
        "JWT. Subsequent calls use that token as Bearer. API version header "
        "2022-11-28."
    )

    d.h(2, "10.2 Webhook verification")
    d.p(
        "X-Hub-Signature-256 must start with sha256=. HMAC-SHA256 of raw body, "
        "compare_digest. Per-org secret is tried first (org looked up by "
        "installation id from payload), then global GITHUB_WEBHOOK_SECRET. "
        "Payload is JSON-parsed before verification in order to find the org — "
        "invalid JSON yields empty dict and likely 401."
    )

    d.h(2, "10.3 Events handled")
    d.bullets([
        "installation / installation_repositories: register repositories_added or repositories; skip existing github_repo_id.",
        "pull_request opened|synchronize: full review pipeline.",
        "Other events: 200 {status: ignored}.",
    ])

    d.h(2, "10.4 Frontend install callback")
    d.p(
        "apps/web/app/api/github/callback/route.ts reads installation_id and state "
        "(org id), requires a Supabase session, POSTs /orgs/github/link, redirects "
        "to /settings/organisation?tab=integrations with connected=true or error=."
    )

    d.h(2, "10.5 Repo sync healing")
    d.p(
        "If the App was installed from github.com without Setup URL, "
        "github_installation_id stays NULL while webhooks still create repos. "
        "_resolve_installation_id uses org field then any non-zero repo.installation_id. "
        "Sync backfills the org field, lists GitHub repos (paginated 100), creates "
        "missing rows, and moves orphans with the same installation_id to this org."
    )

    # ── 11. Sentry ─────────────────────────────────────────────────────────
    d.h(1, "11. Sentry integration")
    d.p(
        "Configure Sentry WebHooks to POST "
        "https://{api-host}/webhooks/sentry?org_id={uuid}. If "
        "SENTRY_WEBHOOK_SECRET is set, header sentry-hook-signature (optional "
        "sha256= prefix) is compared to hex HMAC of the body. Empty secret "
        "disables verification — acceptable only for local experiments."
    )
    d.p(
        "Only payload.action == created is processed. Title and id come from "
        "data.issue; stack from data.event.exception.values[0] last 8 frames. "
        "Claude triage failure uses a P2 fallback. Incident is committed, published "
        "as incident.new, and emailed."
    )

    # ── 12. Realtime ───────────────────────────────────────────────────────
    d.h(1, "12. Real-time system")
    d.p(
        "redis_service.check_redis pings once and caches the boolean. "
        "publish_to_org publishes JSON to channel org:{org_id}:incidents, or on "
        "failure / unavailability calls routers.ws.broadcast_to_org."
    )
    d.p(
        "Each WS connection optionally starts a pubsub listener task that "
        "forwarded Redis messages as text. Client receive loop handles chat and "
        "resolve. Disconnect cancels the listener and unsubscribes. There is no "
        "heartbeat or token refresh on the socket; 5-minute JWT is only checked "
        "at connect time."
    )
    d.p(
        "Frontend: useIncidentWS and useDashboardWS rewrite http→ws on "
        "NEXT_PUBLIC_API_URL (so https becomes wss). Token is placed in the query "
        "string (visible in logs and Referer-adjacent tooling; accepted trade-off "
        "because browsers cannot set WS Authorization headers easily)."
    )

    # ── 13. Email ──────────────────────────────────────────────────────────
    d.h(1, "13. Email and notifications")
    d.p(
        "Resend POST https://api.resend.com/emails with from = RESEND_FROM_ADDRESS "
        "(default onboarding@resend.dev). Incident emails colour-code P1–P4. PR "
        "review emails colour the score (≥80 green, ≥60 orange, else red). "
        "BackgroundTasks run after the HTTP response is scheduled; they receive the "
        "same request-scoped session in some call sites — operators should be aware "
        "that FastAPI BackgroundTasks plus a closed session can fail silently; "
        "current code passes db into send_* functions which query channels on that "
        "session."
    )

    # ── 14. Frontend ───────────────────────────────────────────────────────
    d.h(1, "14. Frontend architecture")

    d.h(2, "14.1 Routing map")
    d.table(
        ["Route group", "Paths", "Purpose"],
        [
            ["Root", "/", "Marketing landing; authed → dashboard"],
            ["Auth", "/login, /sign-up, /auth/callback", "Password + Google"],
            ["App", "/dashboard", "Stats, PR list, incidents cards, WS"],
            ["App", "/dashboard/prs, /dashboard/prs/[id]", "PR inbox and review detail"],
            ["App", "/dashboard/incidents, /incidents/[id]", "Incident list and room"],
            ["App", "/dashboard/my-github", "Personal PRs, branches, stats"],
            ["App", "/dashboard/team", "Team quality + AI grade + weekly report"],
            ["App", "/onboarding", "4-step org / GitHub / Sentry / invite"],
            ["App", "/join", "Accept invitation"],
            ["App", "/profile, /settings, /settings/organisation, /settings/billing", "Profile, GitHub, members, notifications, mock billing"],
            ["Marketing", "/features, /pricing, /docs, /api-reference, /about, /blog, /changelog, /roadmap, /careers, /press, /support, /status, /coming-soon", "Public content"],
        ],
    )

    d.h(2, "14.2 App shell")
    d.p(
        "Root layout wraps AuthProvider, Inter + Playfair fonts, metadata/icons. "
        "(app)/layout is a client component: SWRProvider → OrgProvider → AppNav → "
        "OrgGuard → main max-w-6xl → Footer. OrgGuard redirects to /onboarding if "
        "no org except /onboarding and /dashboard. Spinner only on first load "
        "without cached org so tab state survives token refresh."
    )
    d.p(
        "OrgProvider loads /orgs/mine, prefers localStorage org, else first "
        "membership. Nav tabs: Dashboard, My GitHub, Incident Room, Billing. "
        "Org chip links to organisation settings. User menu: profile, org, billing, "
        "sign out."
    )

    d.h(2, "14.3 Data layer")
    d.p(
        "apiFetch always JSON, Bearer, optional X-Org-Id. SWR keys are "
        "[path, token, orgId] so org switches refetch. POLL_INTERVAL 30s for PRs, "
        "incidents, GitHub activity; 60s for team stats and weekly reports."
    )

    d.h(2, "14.4 Domain types")
    d.p(
        "apps/web/types/index.ts is the canonical frontend contract: Plan, Role, "
        "PRStatus, Severity, IncidentStatus/Severity, Org, Member, Repo, "
        "PullRequest, Incident, TeamStats, WeeklyReport, BranchAssignment, "
        "MyGitHubActivity, WSEvent. Keep API serializers in camelCase aligned with "
        "these interfaces."
    )

    d.h(2, "14.5 UI system")
    d.p(
        "Brand tokens via CSS variables (--ink, --bg, --surface, --card, --border, "
        "--pos, --neg). Marketing uses animated hero, container-scroll, counters. "
        "Product UI is denser cards with serif headings. InteractiveHoverButton is "
        "used on primary CTAs."
    )

    d.h(2, "14.6 Onboarding sequence")
    d.numbered([
        "Create organisation (POST /orgs) and persist org id.",
        "Connect GitHub (install App; deep-link via NEXT_PUBLIC_GITHUB_APP_NAME).",
        "Show Sentry webhook URL containing org id.",
        "Invite teammates (POST /orgs/invite).",
    ])

    d.h(2, "14.7 Security headers (Next)")
    d.p(
        "HSTS max-age=31536000 includeSubDomains preload; X-Content-Type-Options "
        "nosniff; X-Frame-Options DENY; Referrer-Policy strict-origin-when-cross-origin; "
        "Permissions-Policy camera/microphone/geolocation empty; CSP default-src self, "
        "script-src self unsafe-inline (Next inline bootstrap), style-src self "
        "unsafe-inline fonts.googleapis.com, img-src self data https, font-src self "
        "gstatic data, connect-src self + Supabase + API + accounts.google.com, "
        "frame-ancestors none, object-src none, base-uri self, form-action self."
    )

    # ── 15. Security ───────────────────────────────────────────────────────
    d.h(1, "15. Security model")

    d.h(2, "15.1 Controls in place")
    d.bullets([
        "Supabase-managed passwords and OAuth; client-side complexity meter.",
        "Email confirmation required in product policy; API re-checks JWT claims.",
        "HMAC verification for GitHub (required) and Sentry (if secret set).",
        "compare_digest for webhook signatures (timing-safe).",
        "GitHub App JWT 10-minute TTL; WS JWT 5-minute TTL.",
        "Admin RBAC on destructive and integration endpoints.",
        "Rate limits on org create (10/h), invite (30/h), profile (30/h), repo sync (20/h), weekly generate (10/h). Keyed by first X-Forwarded-For hop.",
        "Limiter degrades to memory:// if RATELIMIT_STORAGE_URI is invalid (e.g. unexpanded ${REDIS_URL}) so the process still boots.",
        "API security headers: X-Content-Type-Options, Referrer-Policy.",
        "Frontend CSP / HSTS / frame denial.",
        "Open-redirect guard on auth callback next parameter.",
        "Secrets not returned from GET /orgs/github/config.",
        "PEM repair without logging key material.",
        "CORS allowlist instead of *.",
    ])

    d.h(2, "15.2 Threat notes and residual risk")
    d.bullets([
        "X-Forwarded-For is trusted for rate-limit keys — correct only behind a proxy that overwrites the header. Spoofing can shard limits.",
        "X-Org-Id is attacker-controlled. Authorisation is 'is this JWT a member of the claimed org?' for most routes via Member lookup on mutating admin APIs, but list endpoints that only filter Incident.org_id == header org_id require that membership is enforced in get_verified_org_id. Currently get_verified_org_id does not verify the caller is a Member of x_org_id. A stolen session could set another org's UUID and read that tenant's PRs/incidents if IDs leak. Hardening: resolve org solely from membership tables, not a client header.",
        "GitHub private keys stored as plaintext in Postgres. Prefer envelope encryption or a secret manager before production scale.",
        "Webhook handler parses JSON before HMAC when resolving per-org secrets (classic trade-off). Global secret still required for unknown installations.",
        "Sentry org_id is a query parameter: anyone who knows the UUID and can hit the URL can create incidents if the Sentry secret is unset.",
        "WS token in query string; logs and proxies may retain it for 5 minutes of validity.",
        "SSL context for Postgres disables hostname verification.",
        "script-src 'unsafe-inline' weakens XSS mitigation until Next nonces are adopted.",
        "create_all on startup is convenient but not a substitute for migration review in production.",
        "First-org fallback on GitHub webhooks can attach foreign repos to the wrong tenant.",
        "Billing UI must not be mistaken for a PCI-compliant Stripe integration.",
    ])

    d.h(2, "15.3 Secrets inventory")
    d.table(
        ["Secret", "Where used"],
        [
            ["DATABASE_URL", "asyncpg + Alembic"],
            ["SUPABASE_JWT_SECRET", "HS256 API auth"],
            ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Browser/server Supabase client (public by design)"],
            ["SUPABASE_SERVICE_KEY", "Auth invite API"],
            ["JWT_SECRET", "WebSocket tokens"],
            ["ANTHROPIC_API_KEY", "Claude"],
            ["GITHUB_APP_PRIVATE_KEY / PATH", "App JWT"],
            ["GITHUB_WEBHOOK_SECRET + per-org column", "Webhook HMAC"],
            ["SENTRY_WEBHOOK_SECRET", "Optional HMAC"],
            ["REDIS_URL", "Pub/sub (+ optional rate-limit storage)"],
            ["STRIPE_SECRET_KEY / WEBHOOK_SECRET", "Reserved"],
            ["RESEND_API_KEY", "Email"],
        ],
    )

    # ── 16. Testing ────────────────────────────────────────────────────────
    d.h(1, "16. Testing")
    d.h(2, "16.1 Backend (pytest)")
    d.table(
        ["Module", "Coverage"],
        [
            ["test_auth.py", "Protected routes 401 without/invalid token"],
            ["test_email_verification.py", "Claim combinations; bypass flag"],
            ["test_security.py", "X-Forwarded-For, headers on /health, limiter fallback, 429"],
            ["test_webhooks.py", "GitHub HMAC valid/invalid/tamper/prefix"],
            ["test_claude_service.py", "Mocked client: structured parse, fence strip (does not call Anthropic)"],
            ["test_repo_sync.py", "installation id resolution; missing-repo diff"],
            ["conftest.py", "Env defaults + TestClient after settings patch"],
        ],
    )
    d.p("Run: cd apps/api && pytest. asyncio_mode = auto.")

    d.h(2, "16.2 Frontend (Vitest)")
    d.p("lib/password.test.ts covers validatePassword. npm test in apps/web. No Playwright/Cypress suite in-repo.")

    d.h(2, "16.3 Gaps")
    d.bullets([
        "No live Claude contract tests (eval harness spec).",
        "No integration tests against a real Postgres for org isolation.",
        "No WS protocol tests.",
        "No Sentry webhook payload fixtures in the test list above beyond GitHub HMAC.",
        "Billing and Stripe untested because unimplemented.",
    ])

    # ── 17. Deployment ─────────────────────────────────────────────────────
    d.h(1, "17. Deployment and operations")

    d.h(2, "17.1 Local development")
    d.p("Prerequisites: Node 20+, Python 3.11+, PostgreSQL or Supabase, Redis optional.")
    d.code(
        "cd apps/web && cp .env.example .env.local && npm install && npm run dev\n"
        "cd apps/api && cp .env.example .env && python -m venv .venv\n"
        ".venv\\Scripts\\activate   # Windows\n"
        "pip install -r requirements.txt\n"
        "alembic upgrade head\n"
        "uvicorn main:app --reload"
    )

    d.h(2, "17.2 Frontend — Vercel")
    d.p(
        "vercel.json only sets framework nextjs. Set Root Directory to apps/web. "
        "Required env: NEXT_PUBLIC_SUPABASE_URL, NEXT_PUBLIC_SUPABASE_ANON_KEY, "
        "NEXT_PUBLIC_API_URL (https API host, no trailing slash), "
        "NEXT_PUBLIC_GITHUB_APP_NAME, optionally NEXT_PUBLIC_APP_URL for the "
        "install callback redirect. next.config still mentions "
        "NEXT_PUBLIC_CLERK_PUBLISHABLE_KEY from a historical Clerk migration; "
        "Clerk is not used at runtime."
    )

    d.h(2, "17.3 Backend — Render")
    d.p(
        "render.yaml: python 3.12, pip install -r requirements.txt, "
        "uvicorn main:app --host 0.0.0.0 --port $PORT, healthCheckPath /health, "
        "autoDeploy true, plan free. All application secrets are sync:false — paste "
        "literal values. Render does not interpolate ${VAR} inside values."
    )

    d.h(2, "17.4 Database")
    d.p(
        "Use Supabase session pooler port 5432 with +asyncpg. URL-encode passwords. "
        "main.py logs a warning if :6543/ or pgbouncer appears in DATABASE_URL. "
        "After deploy, run alembic upgrade head from a one-off shell with the same URL."
    )

    d.h(2, "17.5 Keep-warm")
    d.p(
        "Render free sleeps after ~15 minutes idle. cron-job.org GET /health every "
        "~10 minutes keeps the dyno and APScheduler alive so Sunday 23:55 EST weekly "
        "reports can run. Without this, reports only generate on the next request "
        "after wake, or via the admin POST."
    )

    d.h(2, "17.6 Connection matrix")
    d.table(
        ["From", "To", "Protocol / notes"],
        [
            ["Vercel browsers", "Render API", "HTTPS REST; WSS for /ws/*; CORS must allow Vercel origin"],
            ["Vercel SSR / route handlers", "Render API + Supabase", "Auth callback profile upsert; GitHub link"],
            ["Render", "Supabase Postgres", "asyncpg TLS, session pooler"],
            ["Render", "Supabase Auth", "JWT verify; optional /auth/v1/invite"],
            ["Render", "Anthropic", "HTTPS messages.create"],
            ["Render", "GitHub", "App JWT + installation token"],
            ["GitHub", "Render", "Webhooks POST /webhooks/github"],
            ["Sentry", "Render", "POST /webhooks/sentry?org_id="],
            ["Render", "Upstash", "rediss pub/sub"],
            ["Render", "Resend", "HTTPS email"],
            ["cron-job.org", "Render", "GET /health"],
        ],
    )

    # ── 18. Environment ────────────────────────────────────────────────────
    d.h(1, "18. Environment variables")

    d.h(2, "18.1 Backend")
    d.table(
        ["Variable", "Required", "Purpose"],
        [
            ["DATABASE_URL", "Yes", "postgresql+asyncpg://…:5432/postgres"],
            ["REDIS_URL", "No (default localhost)", "redis:// or rediss://"],
            ["SUPABASE_JWT_SECRET", "Yes", "HS256"],
            ["JWT_SECRET", "WS yes", "32+ random chars"],
            ["ANTHROPIC_API_KEY", "Yes", "sk-ant-…"],
            ["GITHUB_APP_ID", "Fallback", "Global App if org has none"],
            ["GITHUB_APP_PRIVATE_KEY", "Prod", "Full PEM paste on Render"],
            ["GITHUB_APP_PRIVATE_KEY_PATH", "Local", "./github-app.pem"],
            ["GITHUB_WEBHOOK_SECRET", "Yes", "Global HMAC fallback"],
            ["SENTRY_WEBHOOK_SECRET", "Recommended", "Empty disables check"],
            ["CORS_ORIGINS", "Prod yes", "Comma-separated origins"],
            ["FRONTEND_URL", "Invites/email", "Default http://localhost:3000"],
            ["SUPABASE_URL / SUPABASE_SERVICE_KEY", "Invites", "Skip email if unset"],
            ["STRIPE_SECRET_KEY / STRIPE_WEBHOOK_SECRET", "Future", "Unused by routers"],
            ["RESEND_API_KEY / RESEND_FROM_ADDRESS", "Email", "Skip send if key empty"],
            ["ENFORCE_EMAIL_VERIFICATION", "No", "Default true"],
            ["RATELIMIT_STORAGE_URI", "No", "Empty = memory; redis:// for multi-worker"],
        ],
    )

    d.h(2, "18.2 Frontend")
    d.table(
        ["Variable", "Purpose"],
        [
            ["NEXT_PUBLIC_API_URL", "REST + derived WS base"],
            ["NEXT_PUBLIC_SUPABASE_URL", "Auth + cookies"],
            ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Public anon key"],
            ["NEXT_PUBLIC_GITHUB_APP_NAME", "Install URL slug"],
            ["NEXT_PUBLIC_APP_URL", "GitHub callback redirect base"],
        ],
    )

    # ── 19. Settings object ────────────────────────────────────────────────
    d.h(1, "19. Configuration object (pydantic Settings)")
    d.p(
        "models/database.py Settings reads .env, extra=ignore. Engine: NullPool, "
        "ssl default context with CERT_NONE, statement_cache_size=0, "
        "database_url.strip(). get_db yields AsyncSessionLocal (expire_on_commit=False)."
    )

    # ── 20. Scheduler ──────────────────────────────────────────────────────
    d.h(1, "20. Weekly report job")
    d.p(
        "Job id weekly_reports, replace_existing True, CronTrigger day_of_week=sun "
        "hour=23 minute=55 timezone America/New_York, callable "
        "services.report_service:run_weekly_reports_for_all_orgs. For each "
        "Organization a new session generates stats (same shape as team-stats, "
        "including parallel branch fetches and Haiku). Failures are logged per org "
        "and do not abort the loop. ImportError of apscheduler/pytz disables the "
        "scheduler with a warning."
    )

    # ── 21. Frontend pages detail ──────────────────────────────────────────
    d.h(1, "21. Product UX behaviour (implementation notes)")
    d.bullets([
        "Dashboard: SWR lists + optional dashboard WS; stats-row (PRs, issues, active incidents, avg MTTR); PR and incident cards.",
        "PR detail: score colour via severityFromScore; comment list by file:line.",
        "Incident room: optimistic local messages only after WS echo; resolve button disabled while resolving until status flips.",
        "Organisation settings: tabs for general, members, GitHub integration (config form, install, repo list, sync), notification channels.",
        "Join page: POST /orgs/join then store org id.",
        "Invitation banner: pending invites from /orgs/my-invites.",
        "Profile: full name upsert + GitHub login for PR attribution.",
    ])

    # ── 22. Historical design decisions ────────────────────────────────────
    d.h(1, "22. Documented design history")
    d.table(
        ["Spec / plan", "Decision"],
        [
            ["2026-05-19 Clerk → Supabase Auth", "Replace Clerk with Supabase Auth; HS256 JWT on API; email/Google."],
            ["2026-06-02 Notification services", "Email channels via Resend; event allowlist; admin CRUD + test send."],
            ["2026-06-14 Free-tier deployment", "Vercel + Render free + Upstash + cron keep-warm; env-driven CORS."],
            ["2026-06-15 Email verification hardening", "Confirm email in Supabase + API enforce flag + CSP/HSTS + rate limits."],
            ["2026-08-04 LLM eval harness", "Do not adopt LangChain; add evals around existing three functions; fix 0–100 score scale first."],
        ],
    )
    d.p(
        "The score-scale bug described in the eval spec (prompt asking 0–10 while "
        "consumers treated 0–100) is corrected in the current claude_service.py "
        "prompt, which explicitly requests an integer 0–100."
    )

    # ── 23. Limitations ────────────────────────────────────────────────────
    d.h(1, "23. Known limitations and technical debt")
    d.numbered([
        "Stripe billing is UI mock + unused columns/keys; plan is not enforced (unlimited usage on free).",
        "PR status merged/closed is in the type system but the webhook does not subscribe to those GitHub actions, so merge_rate in personal stats is usually 0.",
        "Incident.users_affected and error_rate are unused by ingest paths.",
        "WebSocket event type names differ between backend (incident.new) and frontend types (incident.created).",
        "No membership check inside get_verified_org_id for header-supplied org ids.",
        "GitHub secrets at rest are plaintext.",
        "create_all + Alembic dual bootstrapping can confuse empty vs migrated databases.",
        "Single-process WS registry; horizontal scale requires Redis and sticky or shared pub/sub only (already the design).",
        "Claude JSON parse has no schema validation (pydantic model) — extra/missing keys pass through .get defaults.",
        "Diff truncated at 500 lines — large PRs are under-reviewed.",
        "GitHub review is COMMENT only, never REQUEST_CHANGES, so it never blocks merge by itself.",
        "Marketing middleware public-path list may not include all (marketing) routes.",
        "Clerk env remnant in next.config.ts.",
        "debug_token.py / _diag_pem.py are operator debug scripts, not part of the service runtime.",
        "No OpenAPI customisation beyond FastAPI defaults (title DevSentinel API 1.0.0).",
        "Incident org_id has no foreign key — orphan incidents possible if org deleted.",
        "Background email tasks may outlive the request session.",
    ])

    # ── 24. Operational runbook ────────────────────────────────────────────
    d.h(1, "24. Operational runbook (short)")
    d.h(2, "24.1 Health")
    d.p("GET /health must return 200. If Render is sleeping, first request is slow; keep-warm should prevent this.")
    d.h(2, "24.2 Webhook 401")
    d.p("Mismatch of GitHub secret (org vs env), or Sentry secret. Confirm X-Hub-Signature-256 and raw body (no JSON re-encoding).")
    d.h(2, "24.3 PR reviews not appearing")
    d.numbered([
        "Confirm App installed on the repo and installation event registered the repo under the correct org.",
        "Run POST /orgs/github/sync-repos as admin.",
        "Check App ID + PEM (normalize_pem logs nothing on success; GitHub 401 on token mint is logged with body).",
        "Confirm ANTHROPIC_API_KEY and that Claude returned parseable JSON.",
        "Confirm event is pull_request opened/synchronize.",
    ])
    d.h(2, "24.4 Incident Room silent")
    d.p("Check JWT_SECRET, CORS, mixed content (https page to ws://), Redis optional. Token expiry is 5 minutes — reconnect after GET /orgs/ws-token.")
    d.h(2, "24.5 Prepared statement / SSL errors")
    d.p("Switch DATABASE_URL to session pooler :5432. Do not use transaction pooler :6543 with asyncpg.")

    # ── 25. Glossary ───────────────────────────────────────────────────────
    d.h(1, "25. Glossary")
    d.table(
        ["Term", "Meaning in this project"],
        [
            ["Org / organisation", "Tenant workspace; slug unique"],
            ["Member", "User↔org membership with admin|member"],
            ["Installation", "GitHub App install; numeric id"],
            ["Review score", "Claude integer 0–100"],
            ["Incident Room", "Live chat + triage UI for one incident"],
            ["MTTR", "Minutes from created_at to resolved_at"],
            ["P1–P4", "Incident severity (manual create also allows P0)"],
            ["X-Org-Id", "Client-selected tenant header"],
            ["WS token", "Short HS256 JWT distinct from Supabase access token"],
            ["Channel", "Notification destination (email list + events)"],
            ["Weekly report", "Snapshotted TeamStats JSON for a week_of date"],
        ],
    )

    # ── 26. File index ─────────────────────────────────────────────────────
    d.h(1, "26. Primary source index")
    d.table(
        ["Concern", "Primary files"],
        [
            ["App bootstrap", "apps/api/main.py"],
            ["Settings / engine", "apps/api/models/database.py"],
            ["Schema", "apps/api/models/{org,user,pull_request,incident,notification}.py"],
            ["AuthN/Z", "apps/api/middleware/auth.py"],
            ["Rate limit / headers", "apps/api/middleware/security.py"],
            ["GitHub + Sentry ingest", "apps/api/routers/webhooks.py"],
            ["Org domain API", "apps/api/routers/orgs.py"],
            ["Claude", "apps/api/services/claude_service.py"],
            ["GitHub client", "apps/api/services/github_service.py"],
            ["Realtime", "apps/api/services/redis_service.py, routers/ws.py"],
            ["Email", "apps/api/services/email_service.py"],
            ["Weekly job", "apps/api/services/report_service.py"],
            ["Next middleware", "apps/web/middleware.ts"],
            ["API client", "apps/web/lib/api.ts, hooks/use-api.ts"],
            ["Session / org", "apps/web/components/auth/auth-provider.tsx, contexts/org-context.tsx"],
            ["Incident Room", "apps/web/components/incidents/incident-room.tsx, hooks/use-incident-ws.ts"],
            ["CSP", "apps/web/next.config.ts"],
            ["Deploy", "render.yaml, vercel.json"],
            ["Public API catalogue", "apps/web/app/(marketing)/api-reference/endpoints.ts"],
        ],
    )

    d.h(1, "27. Document maintenance")
    d.p(
        "Regenerate or edit this file when adding routers, changing auth, altering "
        "the score contract, or changing deploy topology. The generator script that "
        "produced this Word file is docs/_generate_tech_doc.py; prefer updating "
        "that script and re-running it so tables stay consistent. Date of this "
        f"generation: {date.today().isoformat()}."
    )

    out = Path(__file__).resolve().parent / "DevSentinel-Technical-Documentation.docx"
    d.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    build()
